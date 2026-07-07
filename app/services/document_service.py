from typing import List, Optional
from sqlalchemy import select, delete, text, bindparam
from sqlalchemy.ext.asyncio import AsyncSession
from datetime import datetime

from app.database import async_session
from app.models.document import Document, DocumentChunk
from app.services.embedding_service import EmbeddingService
from app.services.llm_analyzer import LLMAnalyzer
from app.config import settings

# Добавляем импорт pgvector
try:
    from pgvector.sqlalchemy import Vector
except ImportError:
    print("⚠ pgvector не установлен. Установите: pip install pgvector")
    Vector = None


import os
import re


class DocumentService:
    """Сервис для работы с документами и их индексацией."""
    
    # =========================================================================
    # Извлечение текста из файлов
    # =========================================================================
    
    @staticmethod
    def extract_text_from_file(file_path: str, file_ext: str = None) -> str:
        """Извлекает текст из файла в зависимости от расширения."""
        if file_ext is None:
            file_ext = os.path.splitext(file_path)[1].lower()
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        if file_ext == ".pdf":
            return DocumentService._extract_from_pdf(file_path)
        elif file_ext == ".docx":
            return DocumentService._extract_from_docx(file_path)
        elif file_ext == ".txt":
            return DocumentService._extract_from_txt(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {file_ext}")
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Извлекает текст из PDF-файла."""
        try:
            import pypdf
        except ImportError:
            raise ImportError("Библиотека pypdf не установлена. Установите: pip install pypdf")
        
        try:
            text_parts = []
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                if reader.is_encrypted:
                    try:
                        reader.decrypt("")
                    except Exception:
                        raise ValueError(f"PDF зашифрован: {os.path.basename(file_path)}")
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            return DocumentService._clean_extracted_text(text)
            
        except Exception as e:
            if isinstance(e, (ImportError, ValueError)):
                raise
            raise Exception(f"Ошибка чтения PDF {os.path.basename(file_path)}: {e}")
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Извлекает текст из DOCX-файла."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Библиотека python-docx не установлена. Установите: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            text = "\n\n".join(text_parts)
            return DocumentService._clean_extracted_text(text)
            
        except Exception as e:
            if isinstance(e, ImportError):
                raise
            raise Exception(f"Ошибка чтения DOCX {os.path.basename(file_path)}: {e}")
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Извлекает текст из TXT-файла."""
        for encoding in ['utf-8', 'cp1251', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                return DocumentService._clean_extracted_text(text)
            except UnicodeDecodeError:
                continue
        
        raise Exception(f"Не удалось определить кодировку файла {os.path.basename(file_path)}")
    
    @staticmethod
    def _clean_extracted_text(text: str) -> str:
        """Очищает извлечённый текст от мусора."""
        if not text:
            return ""
        
        text = re.sub(r'[\u200B\u200C\u200D\u2060\uFEFF]', '', text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {3,}', '  ', text)
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    # =========================================================================
    # CRUD операции с документами
    # =========================================================================
    
    @staticmethod
    async def create_document(
        tender_id: int,
        doc_type: str,
        title: str,
        content: str,
        email_id: Optional[int] = None,
        metadata: Optional[dict] = None,
    ) -> Document:
        """Создать новый документ."""
        async with async_session() as session:
            doc = Document(
                tender_id=tender_id,
                email_id=email_id,
                doc_type=doc_type,
                title=title,
                content=content,
                doc_metadata=metadata or {},
            )
            session.add(doc)
            await session.commit()
            await session.refresh(doc)
            return doc
    
    @staticmethod
    async def get_document_by_id(document_id: int) -> Optional[Document]:
        """Получить документ по ID."""
        async with async_session() as session:
            result = await session.execute(select(Document).where(Document.id == document_id))
            return result.scalar_one_or_none()
    
    @staticmethod
    async def get_documents_by_tender(tender_id: int) -> List[Document]:
        """Получить все документы тендера."""
        async with async_session() as session:
            result = await session.execute(select(Document).where(Document.tender_id == tender_id))
            return result.scalars().all()
    
    # =========================================================================
    # Индексация документов для RAG
    # =========================================================================
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Разбить текст на чанки с перекрытием."""
        chunk_size = chunk_size or settings.RAG_CHUNK_SIZE
        overlap = overlap or settings.RAG_CHUNK_OVERLAP
        
        if not text or not text.strip():
            return []
        
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(para) > chunk_size:
                sentences = para.replace('. ', '.\n').split('\n')
                for sent in sentences:
                    sent = sent.strip()
                    if not sent:
                        continue
                    if len(current_chunk) + len(sent) < chunk_size:
                        current_chunk += sent + " "
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = sent + " "
            else:
                if len(current_chunk) + len(para) < chunk_size:
                    current_chunk += para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        if overlap > 0 and len(chunks) > 1:
            overlapped_chunks = [chunks[0]]
            for i in range(1, len(chunks)):
                prev_chunk = chunks[i-1]
                overlap_text = prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
                overlapped_chunks.append(overlap_text + " " + chunks[i])
            chunks = overlapped_chunks
        
        return chunks
    
    @staticmethod
    async def index_document(document_id: int) -> int:
        """
        Проиндексировать документ: разбить на чанки, получить эмбеддинги.
        
        Returns:
            Количество созданных чанков
        """
        async with async_session() as session:
            result = await session.execute(select(Document).where(Document.id == document_id))
            doc = result.scalar_one_or_none()
            
            if not doc:
                raise ValueError(f"Document {document_id} not found")
            
            await session.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document_id))
            
            chunks_text = DocumentService.chunk_text(doc.content)
            
            if not chunks_text:
                doc.is_indexed = True
                doc.indexed_at = datetime.utcnow()
                doc.chunks_count = 0
                await session.commit()
                return 0
            
            print(f"  Индексация: {len(chunks_text)} чанков...")
            
            # ⭐ ИСПРАВЛЕНО: используем LLMAnalyzer вместо embedding_service
            llm = LLMAnalyzer()
            
            for i, chunk_text in enumerate(chunks_text):
                embedding = await llm.get_embedding(chunk_text)
                
                chunk = DocumentChunk(
                    document_id=doc.id,
                    tender_id=doc.tender_id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=embedding,
                    chunk_metadata={"chunk_number": i + 1, "total_chunks": len(chunks_text)}
                )
                session.add(chunk)
                
                if (i + 1) % 3 == 0 or i == len(chunks_text) - 1:
                    print(f"    Обработано {i + 1}/{len(chunks_text)} чанков")
            
            doc.is_indexed = True
            doc.indexed_at = datetime.utcnow()
            doc.chunks_count = len(chunks_text)
            
            await session.commit()
            return len(chunks_text)
    
    # =========================================================================
    # Семантический поиск
    # =========================================================================
    
    @staticmethod
    async def search(query: str, tender_id: Optional[int] = None, top_k: int = None) -> List[dict]:
        """Семантический поиск с опциональной фильтрацией по тендеру."""
        top_k = top_k or settings.RAG_TOP_K
        
        async with async_session() as session:
            from sqlalchemy import func
            
            llm = LLMAnalyzer()
            query_embedding = await llm.get_query_embedding(query)
            
            print(f"\n  🔍 === ОТЛАДКА ПОИСКА ===")
            print(f"  🔍 Запрос: '{query}'")
            print(f"  🔍 tender_id: {tender_id}, top_k: {top_k}")
            
            # Проверяем наличие данных
            chunks_count_result = await session.execute(select(func.count(DocumentChunk.id)))
            chunks_count = chunks_count_result.scalar()
            print(f"  🔍 Всего чанков в БД: {chunks_count}")
            
            if chunks_count == 0:
                return []
            
            # ⭐ КЛЮЧЕВОЕ ИЗМЕНЕНИЕ: формируем вектор как литерал
            query_vector_str = "[" + ",".join(str(x) for x in query_embedding) + "]"
            
            # ⭐ Используем raw SQL БЕЗ ORDER BY — сортируем в Python
            if tender_id:
                sql = text(f"""
                    SELECT 
                        dc.id,
                        dc.document_id,
                        dc.tender_id,
                        dc.chunk_index,
                        dc.content,
                        dc.metadata,
                        d.title as document_title,
                        d.doc_type,
                        dc.embedding <=> '{query_vector_str}'::vector as distance
                    FROM document_chunks dc
                    JOIN documents d ON dc.document_id = d.id
                    WHERE dc.tender_id = {int(tender_id)}
                """)
            else:
                sql = text(f"""
                    SELECT 
                        dc.id,
                        dc.document_id,
                        dc.tender_id,
                        dc.chunk_index,
                        dc.content,
                        dc.metadata,
                        d.title as document_title,
                        d.doc_type,
                        dc.embedding <=> '{query_vector_str}'::vector as distance
                    FROM document_chunks dc
                    JOIN documents d ON dc.document_id = d.id
                """)
            
            print(f"  🔍 Выполняем SQL запрос (без ORDER BY)...")
            result = await session.execute(sql)
            rows = result.all()
            
            #print(f"  🔍 SQL вернул строк: {len(rows)}")
            
            # ⭐ Сортируем в Python по distance (меньше = лучше)
            rows_sorted = sorted(rows, key=lambda r: float(r[8]))
            
            # Берём top_k результатов
            rows_top = rows_sorted[:top_k]
            
            sources = []
            for row in rows_top:
                distance = float(row[8])
                similarity = 1.0 - distance
                
                sources.append({
                    "chunk_id": row[0],
                    "document_id": row[1],
                    "tender_id": row[2],
                    "chunk_index": row[3],
                    "content": row[4],
                    "metadata": row[5],
                    "document_title": row[6],
                    "doc_type": row[7],
                    "similarity": similarity
                })
            
            print(f"  🔍 Найдено чанков: {len(sources)}")
            if sources:
                for i, s in enumerate(sources[:3], 1):
                    print(f"    [{i}] similarity={s['similarity']:.4f}, doc={s['document_title']}")
                    print(f"        content: {s['content'][:150]}...")
            print(f"  🔍 === КОНЕЦ ОТЛАДКИ ===\n")
            
            return sources
    
    # =========================================================================
    # Удаление документов
    # =========================================================================
    
    @staticmethod
    async def delete_document(document_id: int) -> bool:
        """Удалить документ и все его чанки."""
        async with async_session() as session:
            result = await session.execute(select(Document).where(Document.id == document_id))
            doc = result.scalar_one_or_none()
            
            if doc and hasattr(doc, 'source_path') and doc.source_path:
                try:
                    if os.path.exists(doc.source_path):
                        os.remove(doc.source_path)
                except Exception as e:
                    print(f"⚠ Не удалось удалить файл {doc.source_path}: {e}")
            
            result = await session.execute(delete(Document).where(Document.id == document_id))
            await session.commit()
            return result.rowcount > 0


# Глобальный экземпляр
document_service = DocumentService()