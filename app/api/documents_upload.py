from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from typing import List
import os
import uuid
from datetime import datetime

from app.database import get_db
from app.models.document import Document
from app.services.embedding_service import EmbeddingService

router = APIRouter(prefix="/api/documents", tags=["documents"])


@router.post("/upload/{tender_id}")
async def upload_documents(
    tender_id: int,
    files: List[UploadFile] = File(...),
    db: AsyncSession = Depends(get_db)
):
    """
    Загрузить файлы для тендера и автоматически индексировать их для RAG.
    """
    uploaded_docs = []
    
    for file in files:
        print(f"\n📎 Загрузка файла: {file.filename}")
        
        # Проверяем расширение
        allowed_extensions = [".pdf", ".docx", ".txt"]
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Неподдерживаемый формат: {file.filename}. Допустимы: {', '.join(allowed_extensions)}"
            )
        
        # Сохраняем файл
        unique_filename = f"{uuid.uuid4()}_{file.filename}"
        upload_dir = "uploads/documents"
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, unique_filename)
        
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        print(f"  ✓ Файл сохранён: {file_path}")
        
        # Извлекаем текст
        extracted_text = ""
        try:
            if file_ext == ".txt":
                extracted_text = content.decode('utf-8', errors='ignore')
            elif file_ext == ".pdf":
                from pypdf import PdfReader
                import io
                reader = PdfReader(io.BytesIO(content))
                extracted_text = "\n".join([p.extract_text() for p in reader.pages[:10]])
            elif file_ext == ".docx":
                from docx import Document
                import io
                doc = Document(io.BytesIO(content))
                extracted_text = "\n".join([p.text for p in doc.paragraphs[:100]])
            elif file_ext == ".doc":
                extracted_text = content.decode('utf-8', errors='ignore')
                extracted_text = ''.join(c for c in extracted_text if c.isprintable() or c in '\n\r\t')
            
            print(f"  ✓ Текст извлечён: {len(extracted_text)} символов")
            
        except Exception as e:
            print(f"  ✗ Ошибка извлечения текста: {e}")
            os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"Ошибка извлечения текста: {str(e)}")
        
        # Создаём запись в БД
        document = Document(
            tender_id=tender_id,
            doc_type="ATTACHMENT",
            title=file.filename,
            content=extracted_text,
            source_path=file_path,
            metadata={
                "original_filename": file.filename,
                "size_bytes": len(content),
                "content_type": file.content_type or "application/octet-stream",
                "uploaded_at": datetime.utcnow().isoformat()
            },
            is_indexed=False,
            created_at=datetime.utcnow()
        )
        
        db.add(document)
        await db.flush()
        
        print(f"  ✓ Документ создан в БД: ID={document.id}")
        
        # ⭐ ВАЖНО: используем EmbeddingService (класс), а не embedding_service (переменную)
        print(f"  🔍 Начинаем индексацию...")
        try:
            result = await EmbeddingService.index_document(db, document.id)
            document.is_indexed = True
            document.indexed_at = datetime.utcnow()
            print(f"  ✓ Документ проиндексирован: {result}")
        except Exception as e:
            print(f"  ✗ Ошибка индексации: {e}")
            import traceback
            traceback.print_exc()
        
        uploaded_docs.append({
            "id": document.id,
            "filename": file.filename,
            "size_bytes": len(content),
            "is_indexed": document.is_indexed
        })
    
    await db.commit()
    
    print(f"\n✓ Загрузка завершена: {len(uploaded_docs)} документ(ов)")
    
    return {
        "status": "success",
        "uploaded_count": len(uploaded_docs),
        "documents": uploaded_docs
    }